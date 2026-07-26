"""DOCX loader (python-docx + Tesseract for embedded images).

A .docx has two kinds of content:
  - typed text (paragraphs, tables) -> read directly, NO OCR needed;
  - embedded images (screenshots, scans) -> OCR each image with Tesseract.
We do NOT convert the document to PDF; that is heavy and unnecessary because the
typed text is already machine-readable. Images live inside the .docx (which is a
zip) under 'word/media/', so we read them straight out and OCR them.
"""
import zipfile
from pathlib import Path

from app.ingestion.loaders.base import DocumentLoader, ExtractedPage, register
from app.ingestion.loaders.ocr import ocr_image_bytes

_IMAGE_EXTS = (".png", ".jpg", ".jpeg", ".bmp", ".tiff", ".gif")


class DocxLoader(DocumentLoader):
    def extract(self, path: Path) -> list[ExtractedPage]:
        import docx

        # 1) typed text (paragraphs + tables)
        document = docx.Document(str(path))
        parts = [p.text for p in document.paragraphs if p.text.strip()]
        for table in document.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))

        # 2) OCR any embedded images (read directly from the .docx zip)
        parts.extend(self._ocr_embedded_images(path))

        text = "\n".join(parts).strip()
        return [ExtractedPage(text=text, metadata={"page": 1})] if text else []

    @staticmethod
    def _ocr_embedded_images(path: Path) -> list[str]:
        results: list[str] = []
        with zipfile.ZipFile(path) as zf:
            for name in zf.namelist():
                if name.startswith("word/media/") and name.lower().endswith(_IMAGE_EXTS):
                    ocr_text = ocr_image_bytes(zf.read(name))
                    if ocr_text:
                        results.append(ocr_text)
        return results


register(".docx", DocxLoader())
